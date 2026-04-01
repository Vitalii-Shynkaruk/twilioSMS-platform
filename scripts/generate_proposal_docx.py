"""Generate SMS Platform Proposal .docx with embedded screenshots."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path
import os

DOCS = Path(__file__).resolve().parent.parent / "docs"
OUT = Path(__file__).resolve().parent.parent / "docs" / "SMS_Platform_Proposal.docx"

SCREENSHOTS = {
    "dashboard": DOCS / "SecureCreditLines-SMS-Platform-03-26-2026_10_09_PM.png",
    "numbers": DOCS / "SecureCreditLines-SMS-Platform-03-26-2026_10_10_PM.png",
    "campaigns": DOCS / "SecureCreditLines-SMS-Platform-03-26-2026_10_11_PM.png",
}

def set_cell_shading(cell, color_hex):
    from docx.oxml import OxmlElement
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading_elm)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2B579A')
    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, 'F2F2F2')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    style_table(table)
    return table


def add_screenshot(doc, key, caption):
    path = SCREENSHOTS.get(key)
    if path and path.exists():
        doc.add_picture(str(path), width=Inches(6.2))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)
    else:
        p = doc.add_paragraph(f"[Screenshot: {caption}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build():
    doc = Document()

    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # ── Title ──
    title = doc.add_heading('SMS Platform — Proven Expertise & Delivery Capability', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Production-proven experience for high-volume Twilio SMS infrastructure\n')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(80, 80, 80)
    run = p.add_run('Live system: https://app.sclcapital.io')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(43, 87, 154)

    doc.add_paragraph()  # spacer

    # ── Overview table ──
    doc.add_heading('Document Structure — 4 Critical Pillars', level=2)

    add_styled_table(doc,
        ['Step', 'Focus Area', 'Why It Matters'],
        [
            ['Step 1', 'A2P 10DLC Registration & Compliance', 'Without proper registration, carriers block unregistered traffic'],
            ['Step 2', 'High-Volume Sending Engine & Infrastructure', 'The core system processing 20,000+ messages/day reliably'],
            ['Step 3', 'Carrier Filtering Mitigation & Deliverability', 'The difference between 60% and 97% delivery in restricted industries'],
            ['Step 4', 'Live Operations, Monitoring & Blast Support', 'Real-time visibility and control during active campaigns'],
        ],
        col_widths=[2, 6, 9]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════
    # STEP 1
    # ══════════════════════════════════════════════
    doc.add_heading('Step 1: A2P 10DLC Registration & Compliance Architecture', level=1)

    doc.add_heading('The Problem', level=3)
    doc.add_paragraph(
        'Sending SMS at scale without proper A2P 10DLC registration means carriers will aggressively filter and block your messages. '
        'In restricted industries (financial services, lending), getting campaign approval is significantly harder — most first submissions get rejected.'
    )

    doc.add_heading('What We Did', level=3)
    doc.add_paragraph(
        'We completed the full A2P 10DLC registration flow for a commercial lending company — a category that carriers treat as high-risk:'
    )

    add_styled_table(doc,
        ['Registration Step', 'Status', 'Details'],
        [
            ['Customer Profile (Twilio Trust Hub)', '✅ Approved', 'Business identity verification — EIN, address, authorized representative'],
            ['Brand Registration (TCR)', '✅ Approved', 'LLC entity registered with The Campaign Registry'],
            ['Campaign Use Case', '❌→✅ Rejected, then Approved', 'First attempt "Low Volume Mixed" rejected. Resubmitted as "Customer Care" — approved'],
            ['Messaging Service', '✅ Linked', 'All 35+ phone numbers attached to A2P-approved Messaging Service'],
        ],
        col_widths=[5, 4, 8]
    )

    doc.add_heading('Key Insight: Why First Submission Failed — And How We Fixed It', level=3)

    p = doc.add_paragraph()
    run = p.add_run('Rejected submission: ')
    run.bold = True
    p.add_run('Use case "Low Volume Mixed" — highest rejection rate across all 10DLC categories for financial services.')

    p = doc.add_paragraph()
    run = p.add_run('Approved resubmission — ')
    run.bold = True
    p.add_run('We changed strategy:')

    items = [
        'Reclassified as "Customer Care" — accurately describes follow-ups, status updates, document requests',
        'Removed all trigger words: "promotional", "offers", "exclusive deal", "limited time"',
        'Emphasized opt-in consent in campaign description (website form, direct request)',
        'All message samples: brand prefix, personalization tokens ({first_name}), opt-out language (Reply STOP)',
        'No URL shorteners (bit.ly, etc.) — only owned domains',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Compliance Engine (Built Into the Platform)', level=3)
    doc.add_paragraph('Every message passes through a real-time compliance gate before sending:')

    add_styled_table(doc,
        ['Check', 'Implementation', 'Latency'],
        [
            ['Suppression List', 'DB lookup + Redis cache (5-min TTL)', '<1ms (cached)'],
            ['Opt-Out Status', 'lead.optedOut flag', 'Real-time'],
            ['Quiet Hours', 'Timezone-aware: 8 PM – 9 AM', 'Real-time'],
            ['Daily Number Limit', 'Per-number cap (350)', 'Real-time'],
            ['Delivery Rate Gate', 'Auto-throttle below 80%', 'Real-time'],
        ],
        col_widths=[4, 8, 4]
    )

    doc.add_heading('Automatic Keyword Compliance', level=3)
    add_styled_table(doc,
        ['Inbound Keyword', 'System Response'],
        [
            ['STOP / CANCEL / UNSUBSCRIBE / END / QUIT', 'Immediately: mark DNC → pause automations → add to suppression list → invalidate cache'],
            ['START / UNSTOP / SUBSCRIBE', 'Re-enable messaging: clear opt-out flag, remove suppression entry'],
            ['HELP / INFO', 'Auto-reply with support information and contact details'],
        ],
        col_widths=[6, 11]
    )

    p = doc.add_paragraph()
    run = p.add_run('Result: ')
    run.bold = True
    p.add_run('Full TCPA compliance with zero impact on sending throughput — all checks cached in Redis, per-message compliance gate runs in microseconds.')

    doc.add_page_break()

    # ══════════════════════════════════════════════
    # STEP 2
    # ══════════════════════════════════════════════
    doc.add_heading('Step 2: High-Volume Sending Engine & Infrastructure', level=1)

    doc.add_heading('The Problem', level=3)
    doc.add_paragraph(
        'Sending 20,000+ messages/day across 35+ numbers requires much more than API calls to Twilio. '
        'Naive implementations create database bottlenecks, fail to handle errors gracefully, and produce inconsistent delivery patterns that trigger carrier filters.'
    )

    doc.add_heading('Architecture', level=3)
    arch_text = (
        'Campaign Request (UI or API)\n'
        '    ↓\n'
        'Audience Selection (tags, status, source filters)\n'
        '    ↓\n'
        'Compliance Pre-Check (batch: suppression + opt-out + quiet hours)\n'
        '    ↓\n'
        'BullMQ Queue (Redis-backed, persistent, crash-safe)\n'
        '    ↓\n'
        'Worker Pool (15 concurrent workers)\n'
        '    ↓\n'
        'Number Selection (health-weighted round-robin)\n'
        '    ↓\n'
        'Twilio API (via A2P Messaging Service)\n'
        '    ↓\n'
        'Delivery Webhook → Status Update → Analytics'
    )
    p = doc.add_paragraph()
    run = p.add_run(arch_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_heading('Throughput Configuration (Production Values)', level=3)
    add_styled_table(doc,
        ['Parameter', 'Value', 'Purpose'],
        [
            ['Per-Number Daily Cap', '350 messages', 'Stay under carrier detection threshold'],
            ['Global Rate Limit', '300 msg/min (5 msg/sec)', 'Twilio API rate compliance'],
            ['Worker Concurrency', '15 parallel workers', 'Optimal for 35+ number pool'],
            ['Queue System', '3 separate BullMQ queues', 'Campaign, Automation, Transactional — isolated priority'],
            ['Retry Strategy', '3 attempts: 5s → 10s → 20s', 'Exponential backoff for transient failures'],
            ['Job Retention', '24h success, 7d failure', 'Full audit trail'],
        ],
        col_widths=[5, 5, 7]
    )

    doc.add_heading('Critical Optimization: Batch Compliance Pre-Loading', level=3)

    p = doc.add_paragraph()
    run = p.add_run('Naive approach (won\'t scale): ')
    run.bold = True
    p.add_run('10,000 leads × 3 DB queries each = 30,000+ individual database queries.')

    p = doc.add_paragraph()
    run = p.add_run('Our approach: ')
    run.bold = True
    p.add_run('Batch-fetch ALL suppressed numbers, opted-out leads, and conversations in 3 queries. '
              'Filter 10,000 leads in-memory. Bulk-create conversations in 1 transaction. Bulk-add jobs in 1 Redis operation. '
              'Total: ~10 queries for any campaign size.')

    p = doc.add_paragraph()
    run = p.add_run('Result: ')
    run.bold = True
    p.add_run('A 10,000-lead campaign loads and starts sending in seconds, not minutes.')

    doc.add_heading('Campaign Features', level=3)
    add_styled_table(doc,
        ['Feature', 'Details'],
        [
            ['Template System', 'Dynamic variables: {{firstName}}, {{company}}, {{lastName}}'],
            ['Spintax', '{Hi|Hey|Hello} {{firstName}} — automatic message variation'],
            ['Audience Filters', 'Tags, status, source, state, date range, explicit lead IDs'],
            ['Sending Speed', 'Configurable 1–600 msg/min with ±40% jitter'],
            ['Safety Gate', 'At least one filter required — prevents accidental "send to all"'],
            ['Scheduling', 'Future-dated campaigns with automatic launch'],
            ['Circuit Breaker', 'Auto-pause at >30% error rate'],
            ['Pause/Resume', 'Instant control — no messages lost, queue preserves state'],
        ],
        col_widths=[4, 13]
    )

    doc.add_heading('Tech Stack', level=3)
    add_styled_table(doc,
        ['Layer', 'Technology'],
        [
            ['Backend', 'Node.js + Express + TypeScript'],
            ['Frontend', 'React 18 + Vite + TailwindCSS'],
            ['Database', 'MySQL 8.0 (Prisma ORM, 18+ tables)'],
            ['Queue', 'Redis 7 + BullMQ (3 queues, 15 workers)'],
            ['Real-Time', 'Socket.IO (WebSocket push)'],
            ['SMS', 'Twilio REST API + Webhooks'],
            ['Auth', 'JWT with refresh token rotation'],
            ['Validation', 'Zod runtime type safety on all endpoints'],
            ['Server', 'Nginx + PM2 + Let\'s Encrypt SSL'],
            ['Hosting', 'DigitalOcean (Ubuntu 24.04)'],
        ],
        col_widths=[4, 13]
    )

    doc.add_paragraph()
    add_screenshot(doc, 'dashboard',
        'Production dashboard: real-time delivery rate (97%), send velocity, error breakdown by Twilio error code, '
        '7-day volume chart — all updating live via WebSocket.')

    doc.add_page_break()

    # ══════════════════════════════════════════════
    # STEP 3
    # ══════════════════════════════════════════════
    doc.add_heading('Step 3: Carrier Filtering Mitigation & Deliverability Optimization', level=1)

    doc.add_heading('The Problem', level=3)
    doc.add_paragraph(
        'In financial services / lending, carrier filtering is the #1 challenge. Carriers actively block messages they classify as spam — '
        'and lending content triggers aggressive filters. Without mitigation, delivery rates can drop to 40–60%.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Our Results: ')
    run.bold = True
    p.add_run('Raised delivery rates from initial ~60% to sustained 83–97% in a restricted industry through 6 layered strategies.')

    doc.add_heading('Strategy 1: 7-Day Number Warm-Up', level=3)
    doc.add_paragraph('New numbers don\'t go to full volume. Automatic ramp prevents carrier suspicion:')
    add_styled_table(doc,
        ['Day', 'Daily Limit', 'Purpose'],
        [
            ['1', '50', 'Establish number reputation'],
            ['2', '100', 'Gradual volume increase'],
            ['3', '150', 'Building carrier trust'],
            ['4', '200', 'Approaching normal capacity'],
            ['5', '250', 'Near-full operation'],
            ['6', '300', 'Almost at capacity'],
            ['7+', '350', 'Full production capacity'],
        ],
        col_widths=[2, 4, 11]
    )

    doc.add_heading('Strategy 2: Health Scoring & Auto-Throttling', level=3)
    doc.add_paragraph('Every number has a real-time health score (0–100). Automatic responses:')
    add_styled_table(doc,
        ['Condition', 'System Action'],
        [
            ['Delivery rate < 80%', 'Capacity reduced to 50%'],
            ['Error streak ≥ 5', 'Number auto-cooled for 24 hours'],
            ['Carrier block error (30007/30034)', 'Immediate cooling + admin alert'],
            ['Campaign error rate > 30%', 'Circuit breaker pauses entire campaign'],
        ],
        col_widths=[6, 11]
    )

    doc.add_heading('Strategy 3: Timing Jitter & Anti-Pattern Detection', level=3)
    items = [
        '±40% random jitter on every message interval',
        'Time distribution — messages spread across business hours window',
        'No burst sending — even "send now" campaigns are queued with natural spacing',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Strategy 4: Content Variation', level=3)
    items = [
        'Spintax: {Hi|Hey|Hello} {{firstName}}, {following up on|checking in about} your {financing|funding} request',
        'Dynamic variables: personalized with recipient name, company, specific details',
        'No trigger words: "promotional", "exclusive offer", "limited time", "act now" — all removed',
        'Conversational tone: messages read like human follow-ups, not marketing blasts',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Strategy 5: Sender Pool Management', level=3)
    add_styled_table(doc,
        ['Pool', 'Purpose', 'Daily Cap/Number', 'Selection Priority'],
        [
            ['Primary', 'Main outbound sending', '350', 'Health-weighted round-robin'],
            ['Warm-Up', 'New numbers in ramp phase', '50–300 (by ramp day)', 'Lower priority, monitored closely'],
            ['Re-engagement', 'Follow-up campaigns', '350', 'Dedicated to reply-based workflows'],
        ],
        col_widths=[3, 5, 4, 5]
    )
    doc.add_paragraph(
        'Sticky sender: Once a lead receives a message from a specific number, all follow-ups come from the same number — builds trust, prevents confusion.'
    )

    doc.add_heading('Strategy 6: Number Lifecycle Management', level=3)
    p = doc.add_paragraph()
    run = p.add_run(
        'PURCHASE → WARMING (7-day ramp) → ACTIVE (full capacity)\n'
        '                                          ↓ (delivery issues)\n'
        '                                       COOLING (24h auto-pause)\n'
        '                                          ↓ (after recovery)\n'
        '                                       ACTIVE (restored)\n'
        '                                          ↓ (persistent failures)\n'
        '                                       SUSPENDED → RETIRED'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_paragraph('All transitions are automatic — the system self-heals without manual intervention.')

    doc.add_paragraph()
    add_screenshot(doc, 'numbers',
        'Number management: A2P registration badge, lifecycle status (WARMING→ACTIVE→COOLING), health score 83/100, '
        'daily usage 247/250, warm-up ramp tracking, pool assignment.')

    doc.add_paragraph()
    add_screenshot(doc, 'campaigns',
        'Campaign results: 15+ campaigns over 8 days in financial services, delivery rates 83–100%, '
        'reply tracking (up to 38 replies per campaign), instant pause/resume control.')

    doc.add_page_break()

    # ══════════════════════════════════════════════
    # STEP 4
    # ══════════════════════════════════════════════
    doc.add_heading('Step 4: Live Operations, Monitoring & Real-Time Blast Support', level=1)

    doc.add_heading('The Problem', level=3)
    doc.add_paragraph(
        'High-volume SMS requires live human oversight during sends. Automated systems catch most issues, '
        'but when delivery drops mid-campaign or a carrier starts blocking, you need instant visibility and instant control.'
    )

    doc.add_heading('Real-Time Monitoring Infrastructure', level=3)

    p = doc.add_paragraph()
    run = p.add_run('WebSocket-powered dashboard (Socket.IO):')
    run.bold = True
    items = [
        'Delivery/failure counters update live during active campaigns — no page refresh',
        'Per-number health visible during blast execution',
        'Error rate trending in real-time',
        'New inbound replies appear instantly in inbox',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('Webhook Processing (Non-Blocking):')
    run.bold = True
    doc.add_paragraph(
        'Twilio status webhooks return 200 OK in <100ms (Twilio never times out). '
        'Actual processing queued via BullMQ with 10 workers: message status update, campaign stats, number health recalculation.'
    )

    doc.add_heading('Error Code Intelligence', level=3)
    add_styled_table(doc,
        ['Twilio Error', 'Meaning', 'Automatic Response'],
        [
            ['30003', 'Unreachable destination', 'Retry with backoff'],
            ['30007', 'Carrier filtering (blocked)', 'Cool number 24h, increment error streak'],
            ['30034', 'Content filter (SMS blocked)', 'Cool number 24h, flag content for review'],
            ['21610', 'Unsubscribed recipient', 'Mark lead DNC, add to suppression'],
            ['21211', 'Invalid phone number', 'Skip, mark lead'],
        ],
        col_widths=[3, 5, 9]
    )

    doc.add_heading('Live Blast Support Capabilities', level=3)
    add_styled_table(doc,
        ['Capability', 'Details'],
        [
            ['Instant Pause/Resume', 'One click — queue preserves all pending messages, zero message loss'],
            ['Mid-Campaign Number Removal', 'Remove underperforming number from rotation without stopping campaign'],
            ['Circuit Breaker', 'Auto-pause at >30% error rate — prevents cascading pool damage'],
            ['Error Code Breakdown', 'Live view: which errors are hitting, which numbers are affected'],
            ['Delivery Rate Trending', 'Real-time chart showing delivery rate climbing or dropping'],
        ],
        col_widths=[5, 12]
    )

    doc.add_heading('What Happens During a Typical 10K Blast', level=3)
    blast = [
        ('T+0:00', 'Campaign launched — 10,000 leads queued'),
        ('T+0:01', 'Compliance pre-check complete (batch: ~10 queries)'),
        ('T+0:02', 'First messages hitting carriers — delivery confirmations streaming in'),
        ('T+0:05', 'Dashboard shows: 450 sent, 420 delivered (93.3%), 0 failures'),
        ('T+0:10', 'Number #7 shows 2 consecutive failures — monitoring but not cooling yet'),
        ('T+0:15', '1,200 sent, 1,130 delivered (94.2%) — healthy campaign'),
        ('T+0:20', 'Number #7 hits 5 errors → auto-cooled. 34 remaining numbers absorb load'),
        ('T+1:30', 'Complete: 9,850 sent, 9,100 delivered (92.4%), 820 replies (8.3%)'),
    ]
    add_styled_table(doc,
        ['Time', 'Event'],
        blast,
        col_widths=[3, 14]
    )

    doc.add_heading('Automation Engine (Drip Sequences)', level=3)
    add_styled_table(doc,
        ['Trigger', 'Example Sequence'],
        [
            ['New Lead Created', 'Welcome message → Day 3 follow-up → Day 7 last chance'],
            ['No Reply After N Days', 'Automated re-engagement with different message angle'],
            ['Status Changed', 'Stage-specific follow-up (qualified → offer presentation)'],
            ['Lead Replies', 'Sequence auto-pauses → human rep takes over conversation'],
        ],
        col_widths=[5, 12]
    )
    doc.add_paragraph(
        'Smart pause: When a lead replies, the automation immediately pauses and hands off to a human rep. '
        'No lead ever gets an automated message while actively engaging.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════
    # SUMMARY
    # ══════════════════════════════════════════════
    doc.add_heading('Summary: Direct Match to Your Requirements', level=1)

    add_styled_table(doc,
        ['Your Requirement', 'Our Proof'],
        [
            ['Twilio SMS infrastructure for high-volume outbound', '✅ Production: 20K+ msg/day, 35+ numbers, BullMQ, 15 workers'],
            ['Compliant strategies for restricted industries', '✅ Built for financial services: TCPA, suppression, quiet hours, keywords'],
            ['A2P 10DLC registration & campaign approvals', '✅ Full flow: rejected → strategy pivot → approved. We know the pitfalls'],
            ['Scaling strategies', '✅ 7-day warm-up, number pools, batch pre-loading for 10K+ campaigns'],
            ['Improve deliverability, reduce carrier filtering', '✅ From ~60% to 83–97% sustained — 6 layered strategies'],
            ['Troubleshoot blocking, filtering, throughput', '✅ Real-time error handling, auto-cooling, circuit breaker, health rotation'],
            ['Number provisioning, messaging services, pools', '✅ Full lifecycle: WARMING→ACTIVE→COOLING, named pools, sticky sender'],
            ['Live support during SMS blasts', '✅ WebSocket dashboard, instant pause/resume, per-number health live'],
            ['Repeatable, scalable system', '✅ Production architecture: campaign → queue → workers → delivery → monitoring'],
            ['Collaborate with leadership', '✅ Clear, actionable communication — strategy + execution, not theory'],
        ],
        col_widths=[7, 10]
    )

    doc.add_paragraph()

    doc.add_heading('What Sets Us Apart', level=2)

    differentiators = [
        ('Actually done this at scale.', 'Not proof-of-concept — 20,000+ messages/day in production, restricted industry, real carrier filtering challenges.'),
        ('Solved the hard problems.', 'A2P rejection → approval. Carrier filtering: 60% → 97%. Auto-healing number pools. Zero-latency compliance at scale.'),
        ('Built the complete system.', 'Lead import → campaign delivery → pipeline management → two-way conversations. Full operational platform.'),
        ('Own both strategy and execution.', 'Architecture, infrastructure, deliverability optimization, live blast support — end-to-end, same team.'),
        ('Production-proven.', 'Platform live at app.sclcapital.io — real data, real compliance, real results.'),
    ]
    for i, (title, desc) in enumerate(differentiators, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {title} ')
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Built and operated in production for a financial services client.\nAvailable for screen-share demonstration of all features described above.')
    run.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

    # ── Save ──
    doc.save(str(OUT))
    print(f"✅ Document saved: {OUT}")


if __name__ == '__main__':
    build()
